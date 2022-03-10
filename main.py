from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import load_workbook
from openpyxl import Workbook
import win32com.client
import PIL
from PIL import ImageGrab, Image
from docx.shared import Inches




location_letters = input("What are the five letters of the filename starting with AA? ")
folder_name = location_letters[:3]
date = input("What was the date of examination? (mm/dd/yyyy) ")

print("Enter the Case Planning File number (usually 1): ")
file_number = input()
document = Document('//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-Case Planning And Treatment Recommendations Worksheet-{}.docx'.format(folder_name, location_letters, file_number))

# Find and save Objective Findings
x = 0
findings = ''
while findings == '':
    paragraph = document.paragraphs[x]
    check1 = paragraph.text
    if check1.find("Symptoms") > 0:
        findings = check1[25:]
    x += 1
    print("Objective Findings: " + findings)


# Find and save Subluxations
x = 0
subluxations = ''
while subluxations == '':
    paragraph = document.paragraphs[x]
    check2 = paragraph.text
    if check2.find("luxations:") > 0:
        subluxations = check2[13:]
    x += 1
    print("Subluxations: " + subluxations)


see_attached = input("See attached (Y/N)?")
# Get treatment phases
"""treatment_phases = [document.paragraphs[12].text, document.paragraphs[22].text, document.paragraphs[32].text]"""
treatment_phases = []
if see_attached != 'y' and see_attached != 'Y':
    x = 1
    y = 0

    all_paras = document.paragraphs
    doc_length = len(all_paras)
    while x < 4:

        while len(treatment_phases) < 3:
            paragraph = document.paragraphs[y]
            check3 = paragraph.text
            if check3.find("per week") > 0:
                treatment_phases.append(document.paragraphs[y].text)
                print(document.paragraphs[y].text)
            print("Treatment Phases: " + str(treatment_phases) + str(len(treatment_phases)))
            if y == doc_length-1:
                break
            y += 1

        x += 1

blood_pressure = input("Blood pressure (XXX/XX): ")

# Positive/Negative tests
test_list = ["Formina Compression", "Jackson's Compression", "Spurling's", "Shoulder Compression", "Adson's Sign", "Nash Sign", "Valsalva", "Scapular Fixation", "Soto-Hall", "Kemp's Sign", "Trendelenburg", "Ely's", "Nachlas", "Goldthwaith's", "SLR, Braggards, WLR, Fajersztajn's", "Romberg's", "Minor's", "Patrick's Faber", "Supported Adams", "Bechiterew", "Gaenslen Buckling"]
test_results = []

print("Positive/Negative tests: Positive left = L, Positive Right = R, LR for Positive-Left and Right, P for Positive. Negative is anything else.")
for count, ele in enumerate(test_list):
    print(ele + " Test: ")
    result = input()
    test_results.append(result.upper())
print(test_results)

# Pain Numbers
enter_list = []
print("Pain Numbers:")
list_description = ["Family/Home Responsibilities: ", "Recreation: ", "Social Activity: ", "Occupation: ", "Self Care: ", "Life Support Activity: "]
for i in range(6):
    category = str(input(list_description[i]))
    enter_list.append(category)

# Range of Motion
print("Range of Motion:")
rom_list = []
rom_description = ["Flexion: ", "Extension: ", "R Lat Flex: ", "L Lat Flex: ", "R Rotation: ", "L Rotation: "]
for x in range(2):
    for i in range(6):
        rom = str(input(rom_description[i]))
        rom_list.append(rom)
print(rom_list)
# Create RoM Chart
workbook = load_workbook('\\\\192.168.0.14\\platinum\\Document\\Template\\RoM for RoF.xlsx')
sheet_1 = workbook['Sheet1']
y=0
for row in range(2, 15):
    if row != 8:
        measured = sheet_1.cell(row, 2)
        temp = rom_list[int(y)]
        measured.value = float(temp)
        y += 1

workbook.save('//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-RoM to RoF-1.xlsx'.format(folder_name, location_letters))

# Create Chart Images
chart_number = 0
input_file = '//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-RoM to RoF-1.xlsx'.format(folder_name, location_letters)
output_image1 = '//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-RoM Chart 1.png'.format(folder_name, location_letters)
output_image2 = '//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-RoM Chart 2.png'.format(folder_name, location_letters)

operation = win32com.client.Dispatch("Excel.Application")
operation.Visible = 0
operation.DisplayAlerts = 0
workbook = operation.Workbooks.Open(input_file)
sheet_1 = operation.Sheets(1)

for x, chart in enumerate(sheet_1.Shapes):
    chart.Copy()
    image = ImageGrab.grabclipboard()
    if x == 0:
        image.save(output_image1, 'png')
    else:
        image.save(output_image2, 'png')
    pass
workbook.Close(True)
operation.Quit()

# Create Report of Findings
#document = Document('//192.168.0.14\\platinum\\Document\\Template\\Report of Findings.DOCX')
document = Document('//192.168.0.14\\platinum\\Document\\Template\\Report Of Findings (Rev. in progress).DOCX')

paragraph = document.paragraphs[12]
run = paragraph.add_run(findings.upper())
font = run.font
font.name = 'Arial'
font.size = Pt(12)
font.bold = True

paragraph = document.paragraphs[14]
run = paragraph.add_run(subluxations.upper())
font = run.font
font.name = 'Arial'
font.size = Pt(12)
font.bold = True

# Test Results
y = 0
for x in range(48, 90, 2):
    paragraph = document.paragraphs[x]
    colon = paragraph.text.find(':')
    whole_paragraph = paragraph.text
    title_part = whole_paragraph[:colon+1]
    after_title = whole_paragraph[colon+1:]

    if test_results[y] == 'L':
        title_part += " (POSITIVE-LEFT)"
    elif test_results[y] == 'R':
        title_part += " (POSITIVE-RIGHT)"
    elif test_results[y] == 'LR':
        title_part += " (POSITIVE-LEFT and RIGHT)"
    elif test_results[y] == 'P':
        title_part += " (POSITIVE)"
    paragraph.text = ''
    run = paragraph.add_run(title_part.upper())
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True
    run = paragraph.add_run(after_title)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)
    y += 1

# BP
paragraph = document.paragraphs[111]
run = paragraph.add_run(blood_pressure)
font = run.font
font.name = 'Arial'
font.size = Pt(12)
font.bold = True

# Table Page
table = document.tables[0]
table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
table.rows[0].cells[0].paragraphs[0].runs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
y = 0
for x in range(1, 14):
    if x != 8:
        table.cell(x, 1).text = rom_list[y]
        y += 1

"""document.add_picture('//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-RoM Chart 1.png'.format(folder_name, location_letters)width=Inches(7), height=Inches(4.21))
document.add_picture('//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-RoM Chart 2.png'.format(folder_name, location_letters)width=Inches(7), height=Inches(4.21))
"""

p = document..paragraphs[118]
r = p.add_run()

r.add_picture('//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-RoM Chart 1.png'.format(folder_name, location_letters), width=Inches(5), height=Inches(3.01))
r.add_picture('//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-RoM Chart 2.png'.format(folder_name, location_letters), width=Inches(5), height=Inches(3.01))
# Last Page
if see_attached.upper() == 'Y' or see_attached.upper() == 'YES':
    paragraph = document.paragraphs[153]
    paragraph.add_run("SEE ATTACHED")
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True

else:
    y = 0
    for x in range(153, 160, 3):
        if len(treatment_phases) == y:
            break
        else:
            paragraph = document.paragraphs[x]
            run = paragraph.add_run(treatment_phases[y].upper())
            font = run.font
            font.name = 'Arial'
            font.size = Pt(12)
            font.bold = True
            y += 1

document.save('//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-Report of Findings-1.docx'.format(folder_name, location_letters))

# Create Questionnaire
document = Document('S:\\REPORTS\\Revisions\\Interim Re-exam Questionnaire.docx')

table = document.tables[1]
table.cell(0, 0).text = findings.upper()
table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True


tablecells = [1, 6, 8, 12, 16, 19]
table = document.tables[2]
for x in range(6):
    table.cell(tablecells[x], 1).text = enter_list[x]
document.save('//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-Interim Re-exam Questionnaire-1.docx'.format(folder_name, location_letters))
print("Questionnaire created")

# Create Interim Report
document = Document('//192.168.0.14\\platinum\\Document\\Template\\Interim Report One Page.DOCX')

table = document.tables[0]
table.cell(0, 0).text = date
table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
table.rows[0].cells[0].paragraphs[0].runs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
y = 0
for x in range(2, 15):
    if x != 8:
        table.cell(x, 1).text = rom_list[y]
        y += 1
document.save('//192.168.0.14\\platinum\\Document\\doc\{}\\{}-Interim Report One Page-1.docx'.format(folder_name, location_letters))
print("Interim Report 1 created")

# Create Final Report
document = Document('//192.168.0.14\\platinum\\Document\\Template\\Final Report With Spreadsheet.DOCX')

table = document.tables[0]
table.cell(0, 0).text = date
table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
table.rows[0].cells[0].paragraphs[0].runs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
y = 0
for x in range(2, 15):
    if x != 8:
        table.cell(x, 1).text = rom_list[y]
        y += 1
document.save('//192.168.0.14\\platinum\\Document\\doc\\{}\\{}-Final Report With Spreadsheet-1.docx'.format(folder_name, location_letters))
print("Interim Report 1 created")

print("Final Report created")
